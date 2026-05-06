#!/usr/bin/env python3
"""Render simple Markdown into DOCX using a reference template."""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Sequence, Union

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError as exc:  # pragma: no cover - dependency guard for CLI users
    raise SystemExit(
        "Missing python-docx. Install it in the active Python environment before "
        "running this helper, or use a pandoc-based conversion path."
    ) from exc


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


@dataclass
class Heading:
    level: int
    text: str


@dataclass
class Paragraph:
    text: str


@dataclass
class BulletList:
    items: list[str]


@dataclass
class NumberList:
    items: list[str]


@dataclass
class TableBlock:
    rows: list[list[str]]


Block = Union[Heading, Paragraph, BulletList, NumberList, TableBlock]


def parse_markdown(md_text: str) -> list[Block]:
    lines = md_text.splitlines()
    blocks: list[Block] = []
    i = 0

    if lines[:1] and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---":
            i += 1
        if i < len(lines):
            i += 1

    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            blocks.append(Heading(level=level, text=stripped[level:].strip()))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines):
            align = lines[i + 1].strip()
            align_chars = align.replace("|", "").replace(":", "").replace("-", "").strip()
            if align.startswith("|") and not align_chars:
                rows: list[list[str]] = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    rows.append([cell.strip() for cell in lines[i].strip().strip("|").split("|")])
                    i += 1
                if len(rows) >= 2:
                    rows = [rows[0]] + rows[2:]
                blocks.append(TableBlock(rows=rows))
                continue

        if re.match(r"^[-*]\s+", stripped):
            items: list[str] = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            blocks.append(BulletList(items=items))
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            blocks.append(NumberList(items=items))
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                i += 1
                break
            if nxt.startswith("#") or nxt.startswith("|") or re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append(Paragraph(text=" ".join(para_lines)))

    return blocks


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def ensure_paragraph_style(doc: Document, style_name: str, base_name: str, font_size_pt: int, bold: bool = False) -> None:
    try:
        style = doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base_name]
    style.font.size = Pt(font_size_pt)
    style.font.bold = bold


def ensure_character_style(doc: Document, style_name: str) -> None:
    try:
        return doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        if style_name == "Inline Code":
            style.element.set(qn("w:styleId"), "InlineCode")
            style.font.name = "Menlo"
        return style


def add_runs(paragraph, text: str) -> None:
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.style = "Inline Code"
        else:
            paragraph.add_run(part)


def paragraph_style_for_heading(level: int, first_h1: bool) -> str:
    if level == 1 and first_h1:
        return "Title"
    if level == 1:
        return "Heading 1"
    if level == 2:
        return "Heading 2"
    if level == 3:
        return "Heading 3"
    return "Heading 4"


def render_blocks(doc: Document, blocks: Sequence[Block], subtitle: str | None = None) -> None:
    ensure_paragraph_style(doc, "Heading 4", "Heading 3", 11, bold=True)
    ensure_character_style(doc, "Inline Code")

    if subtitle:
        try:
            paragraph = doc.add_paragraph(style="Subtitle")
        except KeyError:
            paragraph = doc.add_paragraph()
        paragraph.text = subtitle

    first_h1 = True
    for block in blocks:
        if isinstance(block, Heading):
            style = paragraph_style_for_heading(block.level, first_h1)
            if block.level == 1 and first_h1:
                first_h1 = False
            paragraph = doc.add_paragraph(style=style)
            add_runs(paragraph, block.text)
        elif isinstance(block, Paragraph):
            paragraph = doc.add_paragraph(style="Normal")
            add_runs(paragraph, block.text)
        elif isinstance(block, BulletList):
            for item in block.items:
                try:
                    paragraph = doc.add_paragraph(style="List Bullet")
                except KeyError:
                    paragraph = doc.add_paragraph()
                add_runs(paragraph, item)
        elif isinstance(block, NumberList):
            for index, item in enumerate(block.items, start=1):
                try:
                    paragraph = doc.add_paragraph(style="List Number")
                except KeyError:
                    paragraph = doc.add_paragraph()
                    paragraph.text = f"{index}. "
                add_runs(paragraph, item)
        elif isinstance(block, TableBlock) and block.rows:
            col_count = max(len(row) for row in block.rows)
            table = doc.add_table(rows=len(block.rows), cols=col_count)
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
            for row_index, row in enumerate(block.rows):
                for col_index in range(col_count):
                    cell_text = row[col_index] if col_index < len(row) else ""
                    cell = table.cell(row_index, col_index)
                    cell.text = ""
                    add_runs(cell.paragraphs[0], cell_text)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_markdown", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--subtitle")
    parser.add_argument("--overwrite", action="store_true")
    args = parser.parse_args()

    if args.output_docx.exists() and not args.overwrite:
        raise SystemExit(f"output exists; pass --overwrite to replace: {args.output_docx}")
    if not args.input_markdown.exists():
        raise SystemExit(f"input markdown does not exist: {args.input_markdown}")
    if not args.template.exists():
        raise SystemExit(f"template does not exist: {args.template}")

    blocks = parse_markdown(args.input_markdown.read_text(encoding="utf-8"))
    doc = Document(str(args.template))
    clear_document(doc)
    render_blocks(doc, blocks, subtitle=args.subtitle)
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output_docx))
    print(args.output_docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
